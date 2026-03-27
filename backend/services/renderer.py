import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from models.schema import CVProfile
from templates.config import TEMPLATES, DEFAULT_TEMPLATE
from templates.base import set_font, add_section_heading, add_bullet

# ---------------------------------------------------------------------------
# Language detection (simple heuristic on profile content)
# ---------------------------------------------------------------------------
_FR_WORDS = {"le", "la", "les", "de", "du", "des", "et", "en", "pour", "avec",
             "dans", "sur", "une", "est", "sont", "par", "au", "aux", "qui", "que"}

def _detect_lang(profile: CVProfile) -> str:
    text = (profile.summary + " " + " ".join(
        b for job in profile.experience for b in job.bullets
    )).lower()
    words = re.findall(r"\b\w+\b", text)
    fr_count = sum(1 for w in words if w in _FR_WORDS)
    return "fr" if fr_count >= 5 else "en"

# ---------------------------------------------------------------------------
# Section heading labels
# ---------------------------------------------------------------------------
_LABELS = {
    "en": {
        "summary": "Summary", "experience": "Experience", "education": "Education",
        "skills": "Skills", "projects": "Projects", "certifications": "Certifications",
        "languages": "Languages",
    },
    "fr": {
        "summary": "Résumé", "experience": "Expérience", "education": "Formation",
        "skills": "Compétences", "projects": "Projets", "certifications": "Certifications",
        "languages": "Langues",
    },
}

def _label(key: str, lang: str) -> str:
    return _LABELS.get(lang, _LABELS["en"]).get(key, key.capitalize())


# ---------------------------------------------------------------------------
# Page scaling
# ---------------------------------------------------------------------------
def _estimate_lines(profile: CVProfile) -> int:
    lines = 4
    if profile.summary:
        lines += 1 + max(1, len(profile.summary) // 90)
    for job in profile.experience:
        lines += 3 + len(job.bullets)
    for _ in profile.education:
        lines += 1
    if profile.skills:
        lines += 1 + len(profile.skills) // 2
    for _ in profile.certifications:
        lines += 1
    for proj in (profile.projects or []):
        lines += 2 + max(1, len(proj.description) // 90)
    if profile.languages:
        lines += 1 + len(profile.languages) // 2
    return lines


def _scale(config: dict, profile: CVProfile) -> dict:
    cfg = dict(config)
    lines = _estimate_lines(profile)
    if lines <= 42:
        cfg["space_before_section"] = min(cfg["space_before_section"] + 4, 18)
        cfg["font_size_body"] = min(cfg["font_size_body"] + 0.5, 11.5)
    elif lines >= 80:
        cfg["space_before_section"] = max(cfg["space_before_section"] - 4, 6)
        cfg["font_size_body"] = max(cfg["font_size_body"] - 0.5, 9.0)
        cfg["font_size_name"] = max(cfg["font_size_name"] - 2, 16)
    return cfg


# ---------------------------------------------------------------------------
# Hyperlink helper
# ---------------------------------------------------------------------------
def _add_hyperlink(paragraph, text: str, url: str, font_name: str,
                   font_size: float, color_hex: str):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Color
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_hex)
    rPr.append(color_el)

    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Size
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Skills table — borderless, always even count
# ---------------------------------------------------------------------------
def _render_skills_table(doc: Document, skills: list, font: str,
                         size: float, color: str) -> None:
    if not skills:
        return
    # Ensure even count
    items = list(skills)
    if len(items) % 2 != 0:
        items = items[:-1]
    if not items:
        return

    table = doc.add_table(rows=len(items) // 2, cols=2)
    # Remove table style borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Remove existing tblBorders if any
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            skill = items[row_idx * 2 + col_idx]
            # Clear default paragraph
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"• {skill}")
            set_font(run, font, size, color_hex=color)
            # Remove cell borders
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBorders.append(el)
            tcPr.append(tcBorders)
            # Set cell width to 50%
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), "4680")
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


# ---------------------------------------------------------------------------
# Doc helpers
# ---------------------------------------------------------------------------
def _new_doc(top=0.75, bottom=0.75, left=0.9, right=0.9) -> Document:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left)
        s.right_margin = Inches(right)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)
    return doc


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_contact_para(doc, profile, font, cfg, alignment=None):
    """Build the contact line with hyperlinks for linkedin and github."""
    contact = profile.contact
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(4)

    plain_parts = [x for x in [contact.email, contact.phone, contact.location] if x]
    set_font(p.add_run("  |  ".join(plain_parts)), font, cfg["font_size_contact"],
             color_hex=cfg["color_contact"])

    link_color = cfg.get("color_title", cfg["color_contact"])

    if contact.linkedin:
        set_font(p.add_run("  |  "), font, cfg["font_size_contact"], color_hex=cfg["color_contact"])
        url = contact.linkedin if contact.linkedin.startswith("http") else f"https://{contact.linkedin}"
        _add_hyperlink(p, "LinkedIn", url, font, cfg["font_size_contact"], link_color)

    if contact.github:
        set_font(p.add_run("  |  "), font, cfg["font_size_contact"], color_hex=cfg["color_contact"])
        url = contact.github if contact.github.startswith("http") else f"https://{contact.github}"
        _add_hyperlink(p, "GitHub", url, font, cfg["font_size_contact"], link_color)

    return p


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
def render_cv(profile: CVProfile, template_name: str = DEFAULT_TEMPLATE) -> bytes:
    config = TEMPLATES.get(template_name, TEMPLATES[DEFAULT_TEMPLATE])
    if template_name == "compact":
        return _render_compact(profile, config)
    elif template_name == "modern":
        return _render_modern(profile, config)
    else:
        return _render_classic(profile, config)


def _render_classic(profile: CVProfile, config: dict) -> bytes:
    cfg = _scale(config, profile)
    doc = _new_doc()
    font = cfg["font_name"]
    lang = _detect_lang(profile)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(profile.name), font, cfg["font_size_name"],
             bold=True, color_hex=cfg["color_name"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(profile.title), font, cfg["font_size_title"],
             italic=True, color_hex=cfg["color_title"])

    _build_contact_para(doc, profile, font, cfg, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _render_body(doc, profile, cfg, lang)
    return _to_bytes(doc)


def _render_modern(profile: CVProfile, config: dict) -> bytes:
    cfg = _scale(config, profile)
    doc = _new_doc(left=0.85, right=0.85)
    font = cfg["font_name"]
    lang = _detect_lang(profile)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(profile.name), font, cfg["font_size_name"],
             bold=True, color_hex=cfg["color_name"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(profile.title), font, cfg["font_size_title"],
             color_hex=cfg["color_title"])

    _build_contact_para(doc, profile, font, cfg)
    _render_body(doc, profile, cfg, lang)
    return _to_bytes(doc)


def _render_compact(profile: CVProfile, config: dict) -> bytes:
    cfg = _scale(config, profile)
    doc = _new_doc(top=0.6, bottom=0.6, left=0.75, right=0.75)
    font = cfg["font_name"]
    lang = _detect_lang(profile)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run(profile.name), font, cfg["font_size_name"],
             bold=True, color_hex=cfg["color_name"])

    # Title inline then contact on next line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    set_font(p2.add_run(profile.title), font, cfg["font_size_title"],
             bold=True, color_hex=cfg["color_title"])

    _build_contact_para(doc, profile, font, cfg)
    _render_body(doc, profile, cfg, lang)
    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Shared body
# ---------------------------------------------------------------------------
def _render_body(doc: Document, profile: CVProfile, config: dict, lang: str) -> None:
    font = config["font_name"]
    body_sz = config["font_size_body"]
    body_color = config["color_body"]
    dim_color = config["color_contact"]
    sp_sec = config["space_before_section"]

    if profile.summary:
        add_section_heading(doc, _label("summary", lang), config)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(profile.summary), font, body_sz, color_hex=body_color)

    if profile.experience:
        add_section_heading(doc, _label("experience", lang), config)
        for job in profile.experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(min(sp_sec - 4, 8))
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(job.company), font, body_sz, bold=True, color_hex=body_color)
            set_font(p.add_run(f"  |  {job.start} – {job.end}"), font, body_sz, color_hex=dim_color)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after = Pt(2)
            set_font(p2.add_run(f"{job.title}  ·  {job.location}"),
                     font, body_sz, italic=True, color_hex=body_color)
            for bullet in job.bullets:
                add_bullet(doc, bullet, config)

    if profile.education:
        add_section_heading(doc, _label("education", lang), config)
        for edu in profile.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(edu.institution), font, body_sz, bold=True, color_hex=body_color)
            set_font(p.add_run(f"  |  {edu.degree} in {edu.field}  ·  {edu.year}"),
                     font, body_sz, color_hex=dim_color)

    if profile.skills:
        add_section_heading(doc, _label("skills", lang), config)
        _render_skills_table(doc, profile.skills, font, body_sz, body_color)

    if profile.projects:
        add_section_heading(doc, _label("projects", lang), config)
        for proj in profile.projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(proj.name), font, body_sz, bold=True, color_hex=body_color)
            if proj.year:
                set_font(p.add_run(f"  ·  {proj.year}"), font, body_sz, color_hex=dim_color)
            if proj.description:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(1)
                p2.paragraph_format.space_after = Pt(1)
                set_font(p2.add_run(proj.description), font, body_sz, color_hex=body_color)
            if proj.url:
                p3 = doc.add_paragraph()
                p3.paragraph_format.space_after = Pt(1)
                url = proj.url if proj.url.startswith("http") else f"https://{proj.url}"
                _add_hyperlink(p3, proj.url, url, font, body_sz - 0.5,
                               config.get("color_title", dim_color))

    if profile.certifications:
        add_section_heading(doc, _label("certifications", lang), config)
        for cert in profile.certifications:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(cert.name), font, body_sz, bold=True, color_hex=body_color)
            extras = [x for x in [cert.issuer, cert.year] if x]
            if extras:
                set_font(p.add_run(f"  ·  {' · '.join(extras)}"), font, body_sz, color_hex=dim_color)

    if profile.languages:
        add_section_heading(doc, _label("languages", lang), config)
        _render_skills_table(
            doc,
            [f"{l.language} — {l.level}" for l in profile.languages],
            font, body_sz, body_color
        )
