from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, font_name: str, font_size: int, bold: bool = False, color_hex: str = None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color_hex:
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _hex_to_rgb(hex_str: str):
    return int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)


def _set_paragraph_shading(para, fill_hex: str):
    """Apply a solid background colour to an entire paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_section_heading(doc, text: str, config: dict):
    """Classic/compact underlined section heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(config["space_before_section"])
    para.paragraph_format.space_after = Pt(config["space_after_heading"])
    run = para.add_run(text.upper())
    set_font(
        run,
        config["font_name"],
        config["font_size_section_heading"],
        bold=True,
        color_hex=config["color_section_heading"],
    )

    # Bottom border rule
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def add_modern_section_heading(doc, text: str, config: dict):
    """Modern template: coloured accent rule below heading text."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(config["space_before_section"])
    para.paragraph_format.space_after = Pt(config["space_after_heading"])
    run = para.add_run(text.upper())
    set_font(
        run,
        config["font_name"],
        config["font_size_section_heading"],
        bold=True,
        color_hex=config["color_section_heading"],
    )

    # Coloured bottom border
    accent = config.get("color_accent", config["color_section_heading"])
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), accent)
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def add_modern_header(doc, profile, config: dict):
    """
    Modern template: full-width navy header block containing name, title,
    and contact line rendered in white text.
    Implemented as shaded paragraphs (no table) for ATS safety.
    """
    bg = config.get("color_header_bg", "1F4E79")

    # Padding paragraph (top breathing room)
    pad_top = doc.add_paragraph()
    pad_top.paragraph_format.space_before = Pt(0)
    pad_top.paragraph_format.space_after = Pt(0)
    _set_paragraph_shading(pad_top, bg)
    pad_top.add_run("")

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(4)
    name_para.paragraph_format.space_after = Pt(2)
    _set_paragraph_shading(name_para, bg)
    name_run = name_para.add_run(profile.name)
    set_font(
        name_run,
        config["font_name"],
        config["font_size_name"],
        bold=True,
        color_hex=config["color_name"],
    )

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(2)
    _set_paragraph_shading(title_para, bg)
    title_run = title_para.add_run(profile.title)
    set_font(
        title_run,
        config["font_name"],
        config["font_size_title"],
        color_hex=config["color_contact"],
    )

    # Contact
    contact = profile.contact
    parts = [contact.email, contact.phone, contact.location]
    if contact.linkedin:
        parts.append(contact.linkedin)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(4)
    _set_paragraph_shading(contact_para, bg)
    contact_run = contact_para.add_run("  |  ".join(parts))
    set_font(
        contact_run,
        config["font_name"],
        config["font_size_contact"],
        color_hex=config["color_contact"],
    )

    # Padding paragraph (bottom breathing room)
    pad_bot = doc.add_paragraph()
    pad_bot.paragraph_format.space_before = Pt(0)
    pad_bot.paragraph_format.space_after = Pt(0)
    _set_paragraph_shading(pad_bot, bg)
    pad_bot.add_run("")


def add_bullet(doc, text: str, config: dict):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    set_font(run, config["font_name"], config["font_size_body"])
    return para